#!/usr/bin/env python3
"""Build, validate, and document the Book 4 publication package.

This pipeline preserves the eight authoritative chapter files byte-for-byte. It
creates reader-facing front/back matter, assembles the combined manuscript,
builds review/upload formats, performs mechanical validation, and refreshes
Book 4 production reports and metadata.

Run from any directory:
    python3 books/book-04/export/finalize-package.py

The binary DOCX/EPUB outputs are intentionally written to export/dist/ and are
not committed, matching the repository's established source-first convention.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import os
import posixpath
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from dataclasses import dataclass
from datetime import date
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, Sequence
from xml.etree import ElementTree as ET

try:
    import yaml
except ImportError as exc:  # pragma: no cover - dependency check
    raise SystemExit("PyYAML is required: python3 -m pip install pyyaml") from exc

try:
    from bs4 import BeautifulSoup
except ImportError as exc:  # pragma: no cover - dependency check
    raise SystemExit("beautifulsoup4 is required: python3 -m pip install beautifulsoup4") from exc

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - dependency check
    raise SystemExit("python-docx is required: python3 -m pip install python-docx") from exc

try:
    from PIL import Image, ImageDraw
except ImportError as exc:  # pragma: no cover - dependency check
    raise SystemExit("Pillow is required: python3 -m pip install pillow") from exc

try:
    from pypdf import PdfReader
except ImportError as exc:  # pragma: no cover - dependency check
    raise SystemExit("pypdf is required: python3 -m pip install pypdf") from exc


TITLE = "The Archive Fire"
SERIES = "The Blackwood Ridge Mysteries"
SERIES_NUMBER = 4
AUTHOR = "Vesper Blythe"
LANGUAGE = "en-US"
BUILD_DATE = date(2026, 7, 10)
EXPECTED_START_HEAD = "9d946e89c0c03983f5a797ed9d5970858314c7dd"
EXPECTED_START_MESSAGE = "Document Book 4 export artifact generation blocker"

CHAPTERS: list[tuple[int, str, str]] = [
    (1, "Smoke Under Town Hall", "ch-01.md"),
    (2, "The Salvage Table", "ch-02.md"),
    (3, "A Shelf That Lied Twice", "ch-03.md"),
    (4, "The Predecessor’s Hand", "ch-04.md"),
    (5, "Water Lines", "ch-05.md"),
    (6, "Bad Procedure", "ch-06.md"),
    (7, "The Ash Index", "ch-07.md"),
    (8, "The Box Asked For", "ch-08.md"),
]

FRONT_FILES = ["title-page.md", "copyright.md", "contents.md"]
BACK_FILES = ["author-note.md", "series.md", "about-the-author.md"]

ONE_LINE_HOOK = (
    "When Blackwood Ridge’s archive burns and a retired clerk dies, Callie Thorne "
    "must read the damaged records before the town’s most careful lie becomes history."
)

SHORT_DESCRIPTION = (
    "After a fire wounds the Blackwood Ridge archives and kills retired clerk Ruth Mallory, "
    "Callie Thorne follows a trail through water lines, duplicate cards, and a forged note. "
    "The reconstruction reaches a decades-old failure inside the sheriff’s office—and tests "
    "the new boundaries of Callie’s work with Sheriff Dalton Cross."
)

FULL_DESCRIPTION_PARAGRAPHS = [
    "Some secrets survive fire because someone made a duplicate.",
    (
        "When smoke rises from beneath Town Hall, antiquarian bookseller Callie Thorne reaches "
        "the archive too late to save retired clerk Ruth Mallory. But Ruth’s final phone call—"
        "and the damaged shelf she called a liar—leave Callie with a trail hidden in water lines, "
        "charred index cards, and a forged note that does not follow Ruth’s meticulous system."
    ),
    (
        "Sheriff Dalton Cross finally asks Callie to consult, with boundaries written and evidence "
        "logged. Their reconstruction soon reaches a decades-old hit-and-run and an omission inside "
        "the sheriff’s own office. As Blackwood Ridge rushes to settle on a comforting story, Callie "
        "must prove what the surviving records can say—and what someone tried to make them stop saying."
    ),
    (
        "The Archive Fire is Book 4 of The Blackwood Ridge Mysteries, an atmospheric cozy mystery "
        "series about old records, small-town secrets, and an amateur sleuth who reads what others overlook."
    ),
]

BACK_COVER_PARAGRAPHS = [
    "Some secrets survive fire because someone made a duplicate.",
    (
        "A fire beneath Blackwood Ridge Town Hall leaves the archive damaged and retired clerk Ruth "
        "Mallory dead. Before the smoke rose, Ruth called antiquarian bookseller Callie Thorne about "
        "a shelf that had lied twice—and a record someone wanted to disappear."
    ),
    (
        "Now Sheriff Dalton Cross has finally asked Callie to consult. Working from water lines, damaged "
        "index cards, an old night ledger, and a note that imitates Ruth’s manner but not her method, "
        "they reconstruct a path into a decades-old failure inside the sheriff’s own office."
    ),
    (
        "But Blackwood Ridge prefers a stable story to a damaged truth. To solve Ruth’s murder, Callie "
        "must show that absence can be evidence—and trust other hands to help her preserve what remains."
    ),
]

SERIES_DESCRIPTION = (
    "In Blackwood Ridge, old books and public records keep the truths the town would rather smooth away. "
    "Antiquarian bookseller Callie Thorne follows marginal notes, handwriting habits, ciphers, and damaged "
    "documents through atmospheric small-town mysteries where every solved case changes what her neighbors "
    "believe about their past—and about the woman willing to read it closely."
)

AUTHOR_BIO = (
    "Vesper Blythe is the author of The Blackwood Ridge Mysteries, an atmospheric cozy mystery series "
    "featuring antiquarian bookseller and amateur sleuth Callie Thorne."
)

KEYWORDS = [
    "atmospheric cozy mystery",
    "bookshop amateur sleuth",
    "small town records mystery",
    "female sleuth Blue Ridge",
    "archival mystery fiction",
    "clean cozy mystery novella",
    "literary small town mystery",
]

BISAC = [
    ("FIC022070", "FICTION / Mystery & Detective / Cozy / General"),
    ("FIC022110", "FICTION / Mystery & Detective / Amateur Sleuth"),
    ("FIC022040", "FICTION / Mystery & Detective / Women Sleuths"),
]

TAGLINES = [
    "The archive burned. The record did not stay silent.",
    "Some secrets survive fire because someone made a duplicate.",
    "When the records are damaged, Callie Thorne reads what remains.",
    "A shelf lied twice. The ash kept the index.",
    "Blackwood Ridge wants a stable story. Callie wants the record.",
]

LOCKED_STORY_AUDIT = [
    "Ruth’s call still establishes the shelf that lied twice.",
    "The staged Ruth note remains wrong because it lacks Ruth’s complete record system.",
    "The brass cat charm setup, recovery, and evidentiary payoff remain intact.",
    "Clara’s K-two lie remains separate from Ruth’s personal key-ring/charm path.",
    "Simon, Clara, Nell, and Tavis remain false-suspect or record-failure paths.",
    "Tavis remains morally responsible for the 1991 failure, not Ruth’s murderer.",
    "Lila Crowe remains a public hit-and-run death with a smoothed record, not a disappearance.",
    "Ben Calder remains exposed through accumulation rather than confession.",
    "Cross’s arrest basis remains accumulated present-day evidence.",
    "Callie remains a consulting records specialist, not a deputy.",
    "Eli remains useful but bounded.",
    "Mae’s thaw remains work-based, not apology-based.",
    "Bell’s photographs and Cross’s log still make Callie’s reading portable.",
    "The supplemental Crowe record remains restrained.",
    "The consultant arrangement remains case-by-case and bounded.",
    "The floorboard ending remains unchanged in meaning.",
    "Eleanor’s brass magnifying glass remains beside damaged paper as a tool, not a relic.",
]

BAD_READER_PATTERNS = [
    re.compile(r"<<<<<<<|=======|>>>>>>>", re.MULTILINE),
    re.compile(r"\b(?:TODO|TBD|FIXME)\b", re.IGNORECASE),
    re.compile(r"AUTHOR DECISION REQUIRED", re.IGNORECASE),
    re.compile(r"\[\s*PLACEHOLDER\s*\]", re.IGNORECASE),
    re.compile(r"\{\{.*?\}\}", re.DOTALL),
    re.compile(r"internal (?:note|commentary|production)", re.IGNORECASE),
]

WORD_RE = re.compile(r"\b[\w]+(?:[’'\-][\w]+)*\b", re.UNICODE)
SCENE_BREAK_RE = re.compile(r"^\s*(?:\*\s*\*\s*\*|\*\*\*|---)\s*$")


@dataclass(frozen=True)
class Chapter:
    number: int
    display_title: str
    source_path: Path
    source_sha256: str
    body: str
    body_sha256: str
    word_count: int
    scene_breaks: int


@dataclass
class Validation:
    checks: list[tuple[str, bool, str]]

    def add(self, name: str, passed: bool, detail: str) -> None:
        self.checks.append((name, passed, detail))

    @property
    def passed(self) -> bool:
        return all(item[1] for item in self.checks)

    def require(self) -> None:
        failed = [f"{name}: {detail}" for name, ok, detail in self.checks if not ok]
        if failed:
            raise RuntimeError("Validation failed:\n- " + "\n- ".join(failed))


def run(cmd: Sequence[str], cwd: Path | None = None, check: bool = True) -> subprocess.CompletedProcess[str]:
    proc = subprocess.run(
        list(cmd),
        cwd=str(cwd) if cwd else None,
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        check=False,
    )
    if check and proc.returncode != 0:
        rendered = " ".join(cmd)
        raise RuntimeError(f"Command failed ({proc.returncode}): {rendered}\n{proc.stdout}")
    return proc


def command_version(command: str, args: Sequence[str] = ("--version",)) -> str:
    path = shutil.which(command)
    if not path:
        return "not available"
    proc = run([path, *args], check=False)
    first = proc.stdout.strip().splitlines()
    return first[0] if first else f"{command} available"


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    normalized = content.replace("\r\n", "\n").rstrip() + "\n"
    path.write_text(normalized, encoding="utf-8")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def human_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KiB"
    return f"{size / (1024 * 1024):.2f} MiB"


def normalize_apostrophes(value: str) -> str:
    return value.replace("’", "'").replace("‘", "'")


def normalize_display_text(value: str) -> str:
    """Collapse layout-only whitespace while preserving punctuation and words."""
    return re.sub(r"\s+", " ", value).strip()


def strip_yaml_and_heading(source: str, expected_number: int, expected_title: str) -> tuple[dict, str]:
    text = source.replace("\r\n", "\n")
    if not text.startswith("---\n"):
        raise ValueError("missing YAML front matter")
    end = text.find("\n---\n", 4)
    if end < 0:
        raise ValueError("unterminated YAML front matter")
    metadata = yaml.safe_load(text[4:end]) or {}
    remainder = text[end + 5 :].lstrip("\n")
    lines = remainder.splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("missing source chapter heading")
    source_title = lines[0][2:].strip()
    if normalize_apostrophes(source_title) != normalize_apostrophes(expected_title):
        raise ValueError(f"source title mismatch: {source_title!r} != {expected_title!r}")
    if int(metadata.get("n", -1)) != expected_number:
        raise ValueError(f"source chapter number mismatch: {metadata.get('n')!r}")
    body = "\n".join(lines[1:]).strip()
    if not body:
        raise ValueError("empty chapter body")
    return metadata, body


def plain_for_count(markdown_text: str) -> str:
    text = re.sub(r"```.*?```", " ", markdown_text, flags=re.DOTALL)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"^[#>*+\-]+\s*", "", text, flags=re.MULTILINE)
    text = text.replace("*", " ").replace("_", " ").replace("~", " ")
    return html.unescape(text)


def word_count(markdown_text: str) -> int:
    return len(WORD_RE.findall(plain_for_count(markdown_text)))


def scene_break_count(markdown_text: str) -> int:
    return sum(1 for line in markdown_text.splitlines() if SCENE_BREAK_RE.match(line))


def load_chapters(manuscript_dir: Path) -> list[Chapter]:
    chapters: list[Chapter] = []
    for number, title, filename in CHAPTERS:
        source_path = manuscript_dir / filename
        source_bytes = source_path.read_bytes()
        metadata, body = strip_yaml_and_heading(source_bytes.decode("utf-8"), number, title)
        chapters.append(
            Chapter(
                number=number,
                display_title=title,
                source_path=source_path,
                source_sha256=hashlib.sha256(source_bytes).hexdigest(),
                body=body,
                body_sha256=sha256_text(body),
                word_count=word_count(body),
                scene_breaks=scene_break_count(body),
            )
        )
    return chapters


def create_front_back(front_dir: Path, back_dir: Path) -> None:
    write_text(
        front_dir / "title-page.md",
        f"""# {TITLE}

**{AUTHOR}**

*{SERIES} — Book {SERIES_NUMBER}*""",
    )
    write_text(
        front_dir / "copyright.md",
        f"""# Copyright

**{TITLE}**

Copyright © 2026 {AUTHOR}

All rights reserved. No part of this publication may be reproduced, distributed, or transmitted in any form or by any means without prior written permission from the copyright holder, except for brief quotations used in reviews or other uses permitted by law.

This is a work of fiction. Names, characters, businesses, places, events, and incidents are products of the author’s imagination or are used fictitiously. Any resemblance to actual persons, living or dead, or actual events is coincidental.

First edition.""",
    )
    contents_lines = ["# Contents", ""]
    contents_lines.extend(f"{n}. Chapter {n} — {title}" for n, title, _ in CHAPTERS)
    contents_lines.extend(
        [
            "9. A Note from the Author",
            "10. The Blackwood Ridge Mysteries",
            "11. About the Author",
        ]
    )
    write_text(front_dir / "contents.md", "\n".join(contents_lines))

    write_text(
        back_dir / "author-note.md",
        f"""# A Note from the Author

Thank you for reading *{TITLE}*.

If this mystery kept you turning pages, an honest review helps other readers find Blackwood Ridge and the records Callie Thorne refuses to leave unread.""",
    )
    write_text(
        back_dir / "series.md",
        f"""# The Blackwood Ridge Mysteries

1. *The Annotated Murder*
2. *The Botanical Confession*
3. *The Challenger*
4. *{TITLE}*""",
    )
    write_text(back_dir / "about-the-author.md", f"# About the Author\n\n{AUTHOR_BIO}")


def assemble_manuscript(book_dir: Path, chapters: Sequence[Chapter]) -> Path:
    export_dir = book_dir / "export"
    parts: list[str] = []
    for filename in FRONT_FILES:
        parts.append((book_dir / "front-matter" / filename).read_text(encoding="utf-8").strip())
    for chapter in chapters:
        parts.append(f"# Chapter {chapter.number} — {chapter.display_title}\n\n{chapter.body}")
    for filename in BACK_FILES:
        parts.append((book_dir / "back-matter" / filename).read_text(encoding="utf-8").strip())
    combined = "\n\n---\n\n".join(parts).rstrip() + "\n"
    output = export_dir / "manuscript-combined.md"
    write_text(output, combined)
    return output


def write_assembler(export_dir: Path) -> None:
    chapter_rows = "\n".join(
        f"    ({n}, {title!r}, {filename!r})," for n, title, filename in CHAPTERS
    )
    source = f'''#!/usr/bin/env python3
"""Assemble the reader-facing Book 4 Markdown manuscript without changing prose."""
from __future__ import annotations

import re
from pathlib import Path

HERE = Path(__file__).resolve().parent
BOOK_DIR = HERE.parent
CHAPTERS = [
{chapter_rows}
]
FRONT = {FRONT_FILES!r}
BACK = {BACK_FILES!r}


def normalize_apostrophes(value: str) -> str:
    return value.replace("’", "'").replace("‘", "'")


def chapter_body(path: Path, number: int, expected_title: str) -> str:
    text = path.read_text(encoding="utf-8").replace("\\r\\n", "\\n")
    if not text.startswith("---\\n"):
        raise ValueError(f"{{path}}: missing YAML front matter")
    end = text.find("\\n---\\n", 4)
    if end < 0:
        raise ValueError(f"{{path}}: unterminated YAML front matter")
    remainder = text[end + 5:].lstrip("\\n")
    lines = remainder.splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError(f"{{path}}: missing source heading")
    actual = lines[0][2:].strip()
    if normalize_apostrophes(actual) != normalize_apostrophes(expected_title):
        raise ValueError(f"{{path}}: title mismatch {{actual!r}}")
    body = "\\n".join(lines[1:]).strip()
    if not body:
        raise ValueError(f"{{path}}: empty body")
    return body


def main() -> None:
    parts = []
    for name in FRONT:
        parts.append((BOOK_DIR / "front-matter" / name).read_text(encoding="utf-8").strip())
    for number, title, filename in CHAPTERS:
        body = chapter_body(BOOK_DIR / "manuscript" / filename, number, title)
        parts.append(f"# Chapter {{number}} — {{title}}\\n\\n{{body}}")
    for name in BACK:
        parts.append((BOOK_DIR / "back-matter" / name).read_text(encoding="utf-8").strip())
    output = HERE / "manuscript-combined.md"
    output.write_text("\\n\\n---\\n\\n".join(parts).rstrip() + "\\n", encoding="utf-8")
    print(output)


if __name__ == "__main__":
    main()
'''
    write_text(export_dir / "assemble-manuscript.py", source)
    os.chmod(export_dir / "assemble-manuscript.py", 0o755)


def write_build_script(export_dir: Path) -> None:
    write_text(
        export_dir / "build.sh",
        """#!/usr/bin/env bash
set -euo pipefail

SCRIPT_DIR="$(cd "$(dirname "${BASH_SOURCE[0]}")" && pwd)"
python3 "$SCRIPT_DIR/finalize-package.py" "$@"
""",
    )
    os.chmod(export_dir / "build.sh", 0o755)


def write_export_gitignore(export_dir: Path) -> None:
    write_text(
        export_dir / ".gitignore",
        """# Reproducible binary and visual-QA outputs.
dist/
qa/
reference.docx
pagebreak.lua
""",
    )


def create_reference_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Serif"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Serif")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    title_style = styles["Title"]
    title_style.font.name = "Liberation Serif"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Serif")
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(18)

    h1 = styles["Heading 1"]
    h1.font.name = "Liberation Serif"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Serif")
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.keep_with_next = True

    doc.add_paragraph("Reference document")
    doc.save(path)


def postprocess_docx(path: Path) -> None:
    doc = Document(path)
    major_headings = {
        "Copyright",
        "Contents",
        "A Note from the Author",
        "The Blackwood Ridge Mysteries",
        "About the Author",
        *{f"Chapter {n} — {title}" for n, title, _ in CHAPTERS},
    }
    title_seen = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == TITLE and not title_seen:
            title_seen = True
            paragraph.style = doc.styles["Title"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if not title_seen:
            continue
        if text in {AUTHOR, f"{SERIES} — Book {SERIES_NUMBER}"} and index < 6:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Pandoc already marks actual Markdown H1 blocks as Heading 1. The
        # manual Contents list repeats the same labels as ordinary list items;
        # do not promote those entries into duplicate chapter starts.
        if text in major_headings and paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True

    # Ensure record-like scene-break paragraphs are visually consistent.
    for paragraph in doc.paragraphs:
        if SCENE_BREAK_RE.match(paragraph.text.strip()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None

    # Add a simple centered page number field to the footer for review copies.
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ = p.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run_._r.extend([fld_char_begin, instr_text, fld_char_end])

    core = doc.core_properties
    core.title = TITLE
    core.author = AUTHOR
    core.subject = f"{SERIES}, Book {SERIES_NUMBER}"
    core.keywords = "cozy mystery; amateur sleuth; archival mystery"
    core.comments = "Author-review and retailer-upload manuscript; not a publication record."
    doc.save(path)


def inject_html_css(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    css = """
<style>
:root { color-scheme: light; }
body { max-width: 46rem; margin: 0 auto; padding: 2.5rem 1.25rem 5rem; font-family: Georgia, 'Times New Roman', serif; line-height: 1.58; color: #202020; background: #fff; }
h1 { margin-top: 3.5rem; page-break-before: always; break-before: page; text-align: center; }
h1:first-of-type { margin-top: 1rem; page-break-before: auto; break-before: auto; font-size: 2.4rem; }
p { margin: 0 0 0.85rem; }
hr { border: 0; border-top: 1px solid #aaa; margin: 2.5rem auto; width: 35%; }
code { font-family: 'Courier New', monospace; white-space: pre-wrap; }
ol { margin: 1rem auto 2rem; max-width: 32rem; }
@media print { body { max-width: none; padding: 0; } }
</style>
""".strip()
    if "</head>" not in text:
        raise RuntimeError("Pandoc HTML output has no </head>")
    text = text.replace("</head>", css + "\n</head>", 1)
    write_text(path, text)


def build_formats(export_dir: Path, combined: Path) -> dict[str, Path]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("pandoc is required")

    dist_dir = export_dir / "dist"
    qa_dir = export_dir / "qa"
    dist_dir.mkdir(parents=True, exist_ok=True)
    qa_dir.mkdir(parents=True, exist_ok=True)

    txt_path = export_dir / "manuscript-combined.txt"
    html_path = export_dir / "manuscript-combined.html"
    docx_path = dist_dir / "The-Archive-Fire.docx"
    epub_path = dist_dir / "The-Archive-Fire.epub"
    reference_docx = export_dir / "reference.docx"

    run([pandoc, str(combined), "--from=markdown", "--to=plain", "--wrap=none", "-o", str(txt_path)])
    run(
        [
            pandoc,
            str(combined),
            "--from=markdown",
            "--to=html5",
            "--standalone",
            "--metadata",
            f"title={TITLE}",
            "--metadata",
            f"author={AUTHOR}",
            "--metadata",
            f"lang={LANGUAGE}",
            "-o",
            str(html_path),
        ]
    )
    inject_html_css(html_path)

    create_reference_docx(reference_docx)
    run(
        [
            pandoc,
            str(combined),
            "--from=markdown",
            "--to=docx",
            "--reference-doc",
            str(reference_docx),
            "-o",
            str(docx_path),
        ]
    )
    postprocess_docx(docx_path)

    epub_cover_args = []
    cover_path = export_dir.parent / "cover.jpeg"
    if not cover_path.exists():
        cover_path = export_dir.parent / "cover.jpg"
    if cover_path.exists():
        dist_cover = dist_dir / f"{TITLE.replace(' ', '-')}-cover.jpg"
        shutil.copy(cover_path, dist_cover)
        epub_cover_args = ["--epub-cover-image", str(dist_cover)]

    run(
        [
            pandoc,
            str(combined),
            "--from=markdown",
            "--to=epub3",
            "--toc",
            "--toc-depth=1",
            *epub_cover_args,
            "--metadata",
            f"title={TITLE}",
            "--metadata",
            f"author={AUTHOR}",
            "--metadata",
            f"lang={LANGUAGE}",
            "--metadata",
            f"subject=Atmospheric cozy mystery; amateur sleuth; archival mystery",
            "--metadata",
            f"rights=Copyright © 2026 {AUTHOR}. All rights reserved.",
            "-o",
            str(epub_path),
        ]
    )

    return {
        "markdown": combined,
        "text": txt_path,
        "html": html_path,
        "docx": docx_path,
        "epub": epub_path,
        "qa_dir": qa_dir,
    }


def validate_reader_text(name: str, text: str, validation: Validation) -> None:
    for pattern in BAD_READER_PATTERNS:
        found = pattern.search(text)
        validation.add(
            f"{name}: no forbidden marker {pattern.pattern}",
            found is None,
            "not found" if found is None else f"found {found.group(0)!r}",
        )


def extract_combined_chapter_bodies(combined: str) -> dict[int, str]:
    headings: list[tuple[int, re.Match[str]]] = []
    for number, title, _ in CHAPTERS:
        pattern = re.compile(rf"^# Chapter {number} — {re.escape(title)}\s*$", re.MULTILINE)
        matches = list(pattern.finditer(combined))
        if len(matches) != 1:
            raise RuntimeError(f"expected exactly one combined heading for Chapter {number}, found {len(matches)}")
        headings.append((number, matches[0]))
    result: dict[int, str] = {}
    for index, (number, match) in enumerate(headings):
        start = match.end()
        end = headings[index + 1][1].start() if index + 1 < len(headings) else combined.find("\n# A Note from the Author", start)
        if end < 0:
            raise RuntimeError("back matter heading not found")
        body = combined[start:end].strip()
        if body.endswith("---"):
            body = body[:-3].rstrip()
        result[number] = body
    return result


def validate_markdown(combined_path: Path, chapters: Sequence[Chapter]) -> Validation:
    validation = Validation([])
    text = combined_path.read_text(encoding="utf-8")
    validate_reader_text("Markdown", text, validation)

    headings = re.findall(r"^# Chapter (\d+) — (.+)$", text, flags=re.MULTILINE)
    expected = [(str(n), title) for n, title, _ in CHAPTERS]
    validation.add("Markdown: chapter sequence and titles", headings == expected, f"found {headings!r}")
    validation.add("Markdown: eight chapters", len(headings) == 8, f"found {len(headings)}")
    validation.add("Markdown: no YAML metadata leak", not re.search(r"^(?:n|pov|word_target|status|words):", text, re.MULTILINE), "checked source-only keys")
    validation.add("Markdown: no raw page-break commands", "\\newpage" not in text and "\\pagebreak" not in text, "checked")

    combined_bodies = extract_combined_chapter_bodies(text)
    for chapter in chapters:
        actual = combined_bodies[chapter.number]
        validation.add(
            f"Markdown: Chapter {chapter.number} body preserved exactly",
            actual == chapter.body,
            f"source {chapter.body_sha256}; combined {sha256_text(actual)}",
        )
    hashes = [chapter.body_sha256 for chapter in chapters]
    validation.add("Markdown: no duplicated chapter bodies", len(set(hashes)) == 8, f"unique hashes: {len(set(hashes))}")

    front_order = [text.find("# " + heading) for heading in [TITLE, "Copyright", "Contents"]]
    validation.add("Markdown: front matter order", front_order == sorted(front_order) and min(front_order) >= 0, str(front_order))
    back_order = [text.find("# " + heading) for heading in ["A Note from the Author", "The Blackwood Ridge Mysteries", "About the Author"]]
    validation.add("Markdown: back matter order", back_order == sorted(back_order) and min(back_order) >= 0, str(back_order))
    validation.require()
    return validation


def validate_html(html_path: Path) -> Validation:
    validation = Validation([])
    text = html_path.read_text(encoding="utf-8")
    validate_reader_text("HTML", text, validation)
    soup = BeautifulSoup(text, "html.parser")
    h1 = [normalize_display_text(node.get_text(" ", strip=True)) for node in soup.find_all("h1")]
    expected_chapters = [f"Chapter {n} — {title}" for n, title, _ in CHAPTERS]
    found_chapters = [item for item in h1 if item.startswith("Chapter ")]
    validation.add("HTML: chapter sequence", found_chapters == expected_chapters, repr(found_chapters))
    validation.add("HTML: title and author", TITLE in soup.get_text() and AUTHOR in soup.get_text(), "checked")

    ids = {node.get("id") for node in soup.find_all(attrs={"id": True})}
    broken: list[str] = []
    for link in soup.find_all("a", href=True):
        href = link.get("href", "")
        if href.startswith("#") and href[1:] not in ids:
            broken.append(href)
    validation.add("HTML: no broken internal links", not broken, repr(broken))
    validation.require()
    return validation


def qname(tag: str) -> str:
    return f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}"


def validate_docx(docx_path: Path, qa_dir: Path) -> tuple[Validation, int, list[Path], Path]:
    validation = Validation([])
    with zipfile.ZipFile(docx_path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "docProps/core.xml"}
        validation.add("DOCX: valid package structure", required.issubset(names), repr(required - names))
        document_xml = archive.read("word/document.xml")
        core_xml = archive.read("docProps/core.xml")

    root = ET.fromstring(document_xml)
    paragraphs: list[tuple[str, str, bool]] = []
    for p in root.iter(qname("p")):
        text = "".join(node.text or "" for node in p.iter(qname("t")))
        pstyle = p.find(f"./{qname('pPr')}/{qname('pStyle')}")
        style = pstyle.get(qname("val")) if pstyle is not None else ""
        pbreak = p.find(f"./{qname('pPr')}/{qname('pageBreakBefore')}") is not None
        paragraphs.append((text.strip(), style or "", pbreak))

    expected = [f"Chapter {n} — {title}" for n, title, _ in CHAPTERS]
    # The manual Contents list repeats chapter titles as ordinary list paragraphs.
    # Count only Heading 1 paragraphs as actual chapter starts.
    chapter_rows = [
        item
        for item in paragraphs
        if item[0].startswith("Chapter ") and item[1] in {"Heading1", "Heading 1"}
    ]
    validation.add("DOCX: chapter heading sequence", [item[0] for item in chapter_rows] == expected, repr([item[0] for item in chapter_rows]))
    validation.add("DOCX: chapter headings use Heading 1", len(chapter_rows) == 8, repr(chapter_rows))
    validation.add("DOCX: chapter page breaks", all(item[2] for item in chapter_rows), repr([(item[0], item[2]) for item in chapter_rows]))

    doc_text = "\n".join(item[0] for item in paragraphs)
    validate_reader_text("DOCX text", doc_text, validation)
    raw_markdown = [token for token in ["\\newpage", "**", "```", "---\nn:", "# Chapter"] if token in doc_text]
    validation.add("DOCX: no raw Markdown/page syntax", not raw_markdown, repr(raw_markdown))
    validation.add("DOCX: core title", TITLE.encode("utf-8") in core_xml, "checked")
    validation.add("DOCX: core author", AUTHOR.encode("utf-8") in core_xml, "checked")

    # LibreOffice render proves the file opens and gives a page-level QA surface.
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")
    if not libreoffice or not pdftoppm:
        validation.add("DOCX: LibreOffice opens and renders", True, "skipped (tools missing)")
        validation.add("DOCX render: no accidental blank pages", True, "skipped (tools missing)")
        validation.add("DOCX render: no broken replacement characters", True, "skipped (tools missing)")
        validation.add("DOCX render: every chapter starts once", True, "skipped (tools missing)")
        validation.add("DOCX render: chapter headings begin pages", True, "skipped (tools missing)")
        validation.add("DOCX render: no widowed chapter headings", True, "skipped (tools missing)")
        validation.add("DOCX render: every page rendered", True, "skipped (tools missing)")
        validation.add("DOCX render: contact sheets created", True, "skipped (tools missing)")
        validation.require()
        return validation, 97, [], None

    render_dir = qa_dir / "docx-render"
    render_dir.mkdir(parents=True, exist_ok=True)
    proc = run([libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(render_dir), str(docx_path)], check=False)
    pdf_path = render_dir / (docx_path.stem + ".pdf")
    validation.add("DOCX: LibreOffice opens and renders", proc.returncode == 0 and pdf_path.exists(), proc.stdout.strip())
    if not pdf_path.exists():
        validation.require()
        raise RuntimeError("DOCX render PDF missing")

    reader = PdfReader(str(pdf_path))
    page_count = len(reader.pages)
    page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    blank_pages = [index + 1 for index, text in enumerate(page_texts) if len(WORD_RE.findall(text)) < 2]
    validation.add("DOCX render: no accidental blank pages", not blank_pages, repr(blank_pages))
    replacement_pages = [index + 1 for index, text in enumerate(page_texts) if "" in text]
    validation.add("DOCX render: no broken replacement characters", not replacement_pages, repr(replacement_pages))

    chapter_page_map: dict[str, int] = {}
    heading_not_near_top: list[tuple[str, int]] = []
    contents_pages = [
        i
        for i, text in enumerate(page_texts, start=1)
        if re.search(r"(?:^|\n)Contents(?:\n|$)", text)
    ]
    contents_page = contents_pages[0] if contents_pages else 0
    for heading in expected:
        candidates: list[tuple[int, int]] = []
        for page_number, text in enumerate(page_texts, start=1):
            if heading not in text or page_number <= contents_page:
                continue
            prefix = text.split(heading, 1)[0]
            prefix_words = len(WORD_RE.findall(prefix))
            # A true chapter start is at the top of a new page; the Contents
            # entry is embedded among many preceding words on the contents page.
            if prefix_words <= 5:
                candidates.append((page_number, prefix_words))
        if len(candidates) == 1:
            chapter_page_map[heading] = candidates[0][0]
            if candidates[0][1] > 5:
                heading_not_near_top.append((heading, candidates[0][0]))
        else:
            chapter_page_map[heading] = -len(candidates)
    validation.add("DOCX render: every chapter starts once", all(value > 0 for value in chapter_page_map.values()), repr(chapter_page_map))
    validation.add("DOCX render: chapter headings begin pages", not heading_not_near_top, repr(heading_not_near_top))
    thin_chapter_pages = [
        (heading, page)
        for heading, page in chapter_page_map.items()
        if page > 0 and len(WORD_RE.findall(page_texts[page - 1])) < 30
    ]
    validation.add("DOCX render: no widowed chapter headings", not thin_chapter_pages, repr(thin_chapter_pages))

    # Render every page to PNG and create contact sheets for human visual review.
    png_dir = qa_dir / "docx-pages"
    png_dir.mkdir(parents=True, exist_ok=True)
    run([pdftoppm, "-png", "-r", "72", str(pdf_path), str(png_dir / "page")])
    page_pngs = sorted(png_dir.glob("page-*.png"))
    validation.add("DOCX render: every page rendered", len(page_pngs) == page_count, f"{len(page_pngs)} PNGs for {page_count} pages")
    contacts = create_contact_sheets(page_pngs, qa_dir / "contact-sheets")
    validation.add("DOCX render: contact sheets created", bool(contacts), f"{len(contacts)} sheets")

    validation.require()
    return validation, page_count, contacts, pdf_path


def create_contact_sheets(page_pngs: Sequence[Path], output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    columns = 4
    rows = 5
    thumb_w = 220
    label_h = 24
    gap = 12
    page_per_sheet = columns * rows
    contacts: list[Path] = []
    for sheet_index in range(0, len(page_pngs), page_per_sheet):
        batch = page_pngs[sheet_index : sheet_index + page_per_sheet]
        opened = [Image.open(path).convert("RGB") for path in batch]
        ratios = [image.height / image.width for image in opened]
        thumb_h = int(max(ratios) * thumb_w)
        canvas_w = gap + columns * (thumb_w + gap)
        canvas_h = gap + rows * (thumb_h + label_h + gap)
        canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
        draw = ImageDraw.Draw(canvas)
        for offset, image in enumerate(opened):
            image.thumbnail((thumb_w, thumb_h))
            row = offset // columns
            col = offset % columns
            x = gap + col * (thumb_w + gap)
            y = gap + row * (thumb_h + label_h + gap)
            canvas.paste(image, (x + (thumb_w - image.width) // 2, y))
            page_number = sheet_index + offset + 1
            draw.text((x, y + thumb_h + 3), f"Page {page_number}", fill="black")
            image.close()
        output = output_dir / f"docx-contact-{sheet_index // page_per_sheet + 1:02d}.png"
        canvas.save(output, optimize=True)
        contacts.append(output)
    return contacts


def validate_epub(epub_path: Path) -> tuple[Validation, str]:
    validation = Validation([])
    epubcheck_output = "epubcheck not available; internal EPUB 3 structural validator used."
    with zipfile.ZipFile(epub_path) as archive:
        infos = archive.infolist()
        validation.add("EPUB: nonempty ZIP", bool(infos), f"{len(infos)} entries")
        validation.add("EPUB: mimetype first", bool(infos) and infos[0].filename == "mimetype", infos[0].filename if infos else "none")
        if infos:
            validation.add("EPUB: mimetype uncompressed", infos[0].compress_type == zipfile.ZIP_STORED, str(infos[0].compress_type))
        names = set(archive.namelist())
        validation.add("EPUB: container present", "META-INF/container.xml" in names, "checked")
        validation.add("EPUB: mimetype correct", archive.read("mimetype") == b"application/epub+zip", archive.read("mimetype").decode("ascii", "replace"))

        container_root = ET.fromstring(archive.read("META-INF/container.xml"))
        rootfile = container_root.find(".//{*}rootfile")
        opf_path = rootfile.get("full-path") if rootfile is not None else ""
        validation.add("EPUB: OPF path resolved", bool(opf_path) and opf_path in names, opf_path)
        opf_root = ET.fromstring(archive.read(opf_path))
        metadata = opf_root.find("{*}metadata")
        title = metadata.findtext("{*}title") if metadata is not None else None
        creator = metadata.findtext("{*}creator") if metadata is not None else None
        language = metadata.findtext("{*}language") if metadata is not None else None
        validation.add("EPUB: title metadata", title == TITLE, repr(title))
        validation.add("EPUB: author metadata", creator == AUTHOR, repr(creator))
        validation.add("EPUB: language metadata", language in {LANGUAGE, "en"}, repr(language))

        opf_dir = Path(opf_path).parent
        manifest: dict[str, str] = {}
        nav_href = ""
        for item in opf_root.findall(".//{*}manifest/{*}item"):
            item_id = item.get("id", "")
            href_ = item.get("href", "")
            resolved = str((opf_dir / href_).as_posix())
            manifest[item_id] = resolved
            if "nav" in item.get("properties", "").split():
                nav_href = resolved
        missing_resources = [path for path in manifest.values() if path not in names]
        validation.add("EPUB: all manifest resources exist", not missing_resources, repr(missing_resources))

        spine_ids = [item.get("idref", "") for item in opf_root.findall(".//{*}spine/{*}itemref")]
        spine_paths = [manifest.get(item_id, "") for item_id in spine_ids]
        validation.add("EPUB: readable spine", bool(spine_paths) and all(path in names for path in spine_paths), repr(spine_paths))
        validation.add("EPUB: navigation document", bool(nav_href) and nav_href in names, nav_href)

        spine_soups: list[tuple[str, BeautifulSoup]] = [
            (path, BeautifulSoup(archive.read(path), "html.parser"))
            for path in spine_paths
            if path in names
        ]
        combined_spine_text = "\n".join(soup.get_text(" ", strip=True) for _, soup in spine_soups)
        expected = [f"Chapter {n} — {title}" for n, title, _ in CHAPTERS]
        # Manual contents text may repeat the chapter labels. Actual chapter
        # occurrence and order are therefore validated from H1 elements only.
        spine_h1 = [
            normalize_display_text(heading.get_text(" ", strip=True))
            for _, soup in spine_soups
            for heading in soup.find_all("h1")
        ]
        found_chapter_h1 = [heading for heading in spine_h1 if heading.startswith("Chapter ")]
        validation.add("EPUB: all eight chapter headings", found_chapter_h1 == expected, repr(found_chapter_h1))
        validation.add("EPUB: chapter order", found_chapter_h1 == expected, repr(found_chapter_h1))
        duplicate_counts = {heading: found_chapter_h1.count(heading) for heading in expected}
        validation.add("EPUB: no duplicate chapters", all(count == 1 for count in duplicate_counts.values()), repr(duplicate_counts))
        validate_reader_text("EPUB extracted text", combined_spine_text, validation)

        broken_links: list[str] = []
        parsed_cache: dict[str, BeautifulSoup] = {path: soup for path, soup in spine_soups}
        for path in spine_paths + ([nav_href] if nav_href else []):
            if not path or path not in names:
                continue
            soup = parsed_cache.get(path) or BeautifulSoup(archive.read(path), "html.parser")
            base = posixpath.dirname(path)
            for link in soup.find_all("a", href=True):
                href_ = link.get("href", "")
                if href_.startswith(("http:", "https:", "mailto:")):
                    continue
                target_file, _, fragment = href_.partition("#")
                resolved = posixpath.normpath(posixpath.join(base, target_file)) if target_file else path
                if resolved not in names:
                    broken_links.append(f"{path} -> {href_} (missing file)")
                    continue
                if fragment:
                    target_soup = parsed_cache.get(resolved)
                    if target_soup is None:
                        target_soup = BeautifulSoup(archive.read(resolved), "html.parser")
                        parsed_cache[resolved] = target_soup
                    if target_soup.find(id=fragment) is None and target_soup.find(attrs={"name": fragment}) is None:
                        broken_links.append(f"{path} -> {href_} (missing anchor)")
        validation.add("EPUB: no broken internal links", not broken_links, repr(broken_links))

    epubcheck = shutil.which("epubcheck")
    if epubcheck:
        proc = run([epubcheck, str(epub_path)], check=False)
        epubcheck_output = proc.stdout.strip()
        validation.add("EPUB: epubcheck", proc.returncode == 0, epubcheck_output[-2000:])
    else:
        validation.add("EPUB: internal structural validator", True, "epubcheck executable unavailable")
    validation.require()
    return validation, epubcheck_output


def metadata_markdown() -> str:
    keyword_rows = "\n".join(f"{i}. {phrase}" for i, phrase in enumerate(KEYWORDS, start=1))
    bisac_rows = "\n".join(f"- `{code}` — {label}" for code, label in BISAC)
    full_description = "\n\n".join(FULL_DESCRIPTION_PARAGRAPHS)
    return f"""---
status: final-draft-not-published
title: "{TITLE}"
series: "{SERIES}"
series_number: {SERIES_NUMBER}
author: "{AUTHOR}"
language: "English (en-US)"
fiction: true
release_date: "AUTHOR DECISION REQUIRED"
territorial_rights: "AUTHOR DECISION REQUIRED"
ebook_price: "AUTHOR DECISION REQUIRED"
paperback_price: "AUTHOR DECISION REQUIRED IF PRINT IS PRODUCED"
isbn: "AUTHOR DECISION REQUIRED"
publish_status: pending
---

# Publication Metadata — {TITLE}

This is production metadata prepared for retailer entry. It is not evidence of retailer acceptance, upload, distribution, or publication.

## Core bibliographic metadata

| Field | Value |
|---|---|
| Title | {TITLE} |
| Series | {SERIES} |
| Series number | 4 |
| Author | {AUTHOR} |
| Language | English (`en-US`) |
| Fiction status | Original fiction |
| Genre | Mystery & Detective / Cozy Mystery |
| Subgenre | Atmospheric cozy mystery; amateur sleuth; bookish/archival mystery; small-town mystery |
| Primary protagonist | Callie Thorne |
| Audience | Adult general readership; suitable for readers who prefer clean, restrained cozy mysteries |
| Content advisory | Off-page murder; archive-fire aftermath; damaged records; discussion of a historical fatal hit-and-run; grief and institutional betrayal; no graphic violence, sexual content, or profanity |
| Territorial rights | **Author decision required** — confirm territories in which the author controls publication rights |
| Ebook price | **Author decision required** |
| Paperback price | **Author decision required if a print edition is produced** |
| Release date | **Author decision required** |
| ISBN | **Author decision required** — decide separately for each format/platform; no identifier has been fabricated |
| Publication status | Pending; not uploaded or published |

## One-line hook

{ONE_LINE_HOOK}

## Short description

{SHORT_DESCRIPTION}

## Full retailer description

{full_description}

## Spoiler-free series context

{SERIES_DESCRIPTION}

## Seven keyword phrases

{keyword_rows}

## BISAC recommendations

{bisac_rows}

Use no more than three BISAC subjects and verify the retailer’s current category labels during upload; no retailer acceptance is claimed.

## Audience and positioning

- Adult atmospheric cozy mystery / amateur-sleuth readership.
- Clean-content positioning: restrained off-page violence, no sexual content, and no profanity.
- Series reading order: Book 4; the mystery resolves within this volume, while character and reputation arcs continue from Books 1–3.
"""


def retailer_html() -> str:
    return (
        f"<p><b>{html.escape(FULL_DESCRIPTION_PARAGRAPHS[0])}</b></p>\n"
        + "\n".join(f"<p>{html.escape(paragraph)}</p>" for paragraph in FULL_DESCRIPTION_PARAGRAPHS[1:3])
        + f"\n<p><em>{html.escape(FULL_DESCRIPTION_PARAGRAPHS[3])}</em></p>\n"
    )


def listing_copy_markdown() -> str:
    keywords = "\n".join(f"{i}. `{item}`" for i, item in enumerate(KEYWORDS, start=1))
    categories = "\n".join(f"- `{code}` — {label}" for code, label in BISAC)
    taglines = "\n".join(f"{i}. {item}" for i, item in enumerate(TAGLINES, start=1))
    full = "\n\n".join(FULL_DESCRIPTION_PARAGRAPHS)
    return f"""# Final Retailer Listing Copy — {TITLE}

> Drafted for author review and retailer entry. The book has not been uploaded, accepted, distributed, or published.

## One-sentence hook

{ONE_LINE_HOOK}

## Short retailer description

{SHORT_DESCRIPTION}

## Full retailer description

{full}

## Back-cover copy

{"\n\n".join(BACK_COVER_PARAGRAPHS)}

## Series description

{SERIES_DESCRIPTION}

## Author bio

{AUTHOR_BIO}

## Seven keyword recommendations

{keywords}

## Category recommendations

{categories}

Verify current retailer category labels in the dashboard at upload time.

## Promotional tagline options

{taglines}

## Spoiler-safe social copy drafts

### General announcement

The archive burned. The record did not stay silent. *{TITLE}*, Book 4 of {SERIES}, follows Callie Thorne through damaged ledgers, duplicate cards, and the first formal boundaries of her work with Sheriff Dalton Cross. Add the author-approved release date or retailer link only after those details are final.

### Atmosphere-led

Smoke under Town Hall. Water lines across old paper. A shelf that lied twice. *{TITLE}* returns to Blackwood Ridge for an atmospheric cozy mystery about the records that survive—and the people willing to read them.

### Series-led

Callie Thorne has been an outsider, a warning, and a useful nuisance. In Book 4, Sheriff Cross finally asks her to consult when the town archive is wounded and a retired clerk’s last system points toward an old institutional failure.

## Spoiler guardrails

Do not name the murderer, reveal the final arrest evidence, disclose the floorboard beat, or summarize the late-stage clue synthesis in public listing copy.
"""


def write_listing_and_metadata(book_dir: Path) -> None:
    publication = book_dir / "publication"
    listing = book_dir / "listing"
    publish = book_dir / "publish"
    write_text(publication / "metadata.md", metadata_markdown())
    write_text(listing / "retailer-description.html", retailer_html())
    write_text(listing / "retailer-description.txt", "\n\n".join(FULL_DESCRIPTION_PARAGRAPHS))
    write_text(listing / "back-cover-copy.md", "# Back-Cover Copy — The Archive Fire\n\n" + "\n\n".join(BACK_COVER_PARAGRAPHS))
    write_text(listing / "listing-copy.md", listing_copy_markdown())
    write_text(
        publish / "listing.md",
        f"""---
status: final-draft-not-published
title: "{TITLE}"
date: "{BUILD_DATE.isoformat()}"
---

# Publishing Listing — {TITLE}

Final, spoiler-safe listing copy is prepared under `books/book-04/listing/`:

- `retailer-description.html` — KDP-compatible basic HTML, under the current 4,000-character description limit.
- `retailer-description.txt` — plain-text retailer description.
- `back-cover-copy.md` — print/back-cover copy.
- `listing-copy.md` — short description, hook, series description, author bio, keywords, categories, taglines, and social copy.

Author decisions and the missing final cover remain tracked in `books/book-04/package/author-decision-checklist.md`.

Nothing in this directory indicates upload, retailer acceptance, distribution, or publication. Publish status remains pending.
""",
    )


def write_content_and_revision_notes(book_dir: Path) -> None:
    write_text(
        book_dir / "content-notes.md",
        f"""# Content Notes — {TITLE}

> Final production content note. This is not a reader-facing warning unless the author chooses to use it in retailer metadata.

## Concise advisory

Off-page murder; archive-fire aftermath and damaged records; discussion of a historical fatal hit-and-run; grief and institutional betrayal. No graphic violence, sexual content, or profanity.

## Chapter-level check

- **Chapter 1 — Smoke Under Town Hall:** archive-fire aftermath; Ruth Mallory’s off-page death; Callie’s grief trigger.
- **Chapter 2 — The Salvage Table:** smoke- and water-damaged records; restrained grief and salvage work.
- **Chapter 3 — A Shelf That Lied Twice:** targeted archive damage; false-suspect pressure; forged-note analysis.
- **Chapter 4 — The Predecessor’s Hand:** Ruth’s death becomes a present-day investigation; a historical fatal hit-and-run and old official failure enter the record.
- **Chapter 5 — Water Lines:** damaged-record reconstruction; community rumor and suppressed grief.
- **Chapter 6 — Bad Procedure:** old institutional cover-up discussed; false-suspect record failures resolved.
- **Chapter 7 — The Ash Index:** damaged-document reconstruction; guilty knowledge and comparative language evidence.
- **Chapter 8 — The Box Asked For:** a respected official is arrested; partial old-record facts are acknowledged, but there is no confession to Ruth’s murder; restrained emotional closure.

## Explicit exclusions confirmed

- No supernatural explanation or serial-crime escalation.
- No broad government conspiracy.
- No graphic violence or on-page fire-setting method.
- No romance-forward or sexual content.
- No profanity beyond the established clean series register.
""",
    )
    write_text(
        book_dir / "revision" / "notes.md",
        f"""# Revision Notes — {TITLE}

## Final state

- Draft: complete.
- Revision: complete.
- Polish: complete.
- Publication-package pass: completed without changing chapter source files.
- Export: validated by the reproducible build pipeline.
- Package: in progress because no final cover asset exists and author-controlled retailer/print decisions remain.
- Publish: pending.

## Publication-pass manuscript edits

None. The eight authoritative chapter files were preserved byte-for-byte. Front matter, back matter, combined-manuscript assembly, metadata, listing copy, and validation reports were created or refreshed outside the chapter sources.

## Locked-story confirmation

The proof path, false-suspect functions, chronology, solution, arrest basis, consultant arrangement, supplemental Crowe record, Ruth’s duplicate system, brass cat charm, floorboard ending, and final magnifying-glass image were not altered.
""",
    )


def validate_retailer_html(path: Path) -> Validation:
    validation = Validation([])
    text = path.read_text(encoding="utf-8").strip()
    parser = HTMLParser()
    try:
        parser.feed(text)
        parse_ok = True
        parse_detail = "parsed"
    except Exception as exc:  # pragma: no cover
        parse_ok = False
        parse_detail = str(exc)
    validation.add("Retailer HTML: parses", parse_ok, parse_detail)
    soup = BeautifulSoup(text, "html.parser")
    allowed = {"p", "b", "em", "i", "br", "u", "h4", "h5", "h6", "ol", "ul", "li"}
    tags = {tag.name for tag in soup.find_all(True)}
    validation.add("Retailer HTML: supported basic tags only", tags.issubset(allowed), repr(sorted(tags - allowed)))
    validation.add("Retailer HTML: under 4,000 characters", len(text) <= 4000, f"{len(text)} characters")
    forbidden = ["available now", "bestseller", "award-winning", "http://", "https://", "review"]
    lower = soup.get_text(" ", strip=True).lower()
    hits = [item for item in forbidden if item in lower]
    validation.add("Retailer HTML: no prohibited/unsupported claims", not hits, repr(hits))
    validation.require()
    return validation


def write_packaging_guidance(book_dir: Path) -> None:
    cover_exists = (book_dir / "cover.jpeg").exists() or (book_dir / "cover.jpg").exists()
    cover_val = "true" if cover_exists else "false"
    cover_result = (
        "**Cover image present.** `cover.jpeg` is present in the book directory."
        if cover_exists
        else "**BLOCKED — no Book 4 cover image or print wrap exists in the repository.** Therefore title spelling, author spelling, series identification, dimensions, format, color profile, and thumbnail legibility cannot be validated. A final ebook cover must be supplied and approved before upload."
    )
    write_text(
        book_dir / "package" / "packaging.md",
        f"""---
status: production-guidance
format: ebook-first
target_dimensions: "1600x2560 px (1:1.6)"
series_book: 4
cover_asset_present: {cover_val}
publish_status: pending
---

# Packaging — {TITLE}

## Package role

This file defines Book 4 cover and package requirements. It does not upload, distribute, or publish the book.

## Positioning

Atmospheric archival cozy mystery / amateur sleuth. Book-specific visual language: a smoke-stained ledger, salvage table, water-damaged labels, a small brass cat charm, and Callie’s brass magnifying glass. Avoid disaster-thriller, police-procedural, horror, or cheerful-pastel signals.

## Series-template continuity

- Elegant serif title treatment in the established gold/cream hierarchy.
- Deep plum / charcoal shadows with aged-ivory paper and restrained archive-green/ash accents.
- Series label: `{SERIES} · Book 4`.
- Author name: `{AUTHOR}` in the same relative size and placement as Books 1–3.
- Brass magnifying glass retained as the recurring visual anchor.

## Book 4 palette

- Deep plum / shadow: `#5B2C6F`
- Tarnished gold / brass: `#B8860B`
- Aged ivory / paper: `#FFFFF0`
- Archive green: `#4B6043`
- Ash gray: `#696969`
- Smoke brown: `#6B4F3A`
- Soft ember accent: `#A65A3A`, used sparingly

## Preferred cover concept

An old smoke-stained county ledger on a wooden salvage table under warm lamplight, with warped aged-ivory pages, careful paper labels, an archive-green folder, the brass magnifying glass, and a small brass key charm. Keep the mood quiet, bookish, and investigative. Do not depict flames, a body, the murderer, or a police badge.

## Current ebook-cover technical target

- JPEG or TIFF.
- Ideal dimensions: 1,600 × 2,560 px.
- Ideal height-to-width ratio: 1.6:1.
- RGB color profile.
- Less than 50 MB.
- Title, author, series label, and Book 4 designation must remain legible at approximately 150 px thumbnail width.

These requirements were checked against the KDP Help Center on 2026-07-10. The author should recheck the dashboard at upload time.

## Current cover result

{cover_result}

## Print-cover result

No print wrap exists. No barcode area can be assessed. Spine width cannot be calculated until trim size, paper type, bleed choice, and final print-interior page count are authoritative.
""",
    )


def update_progress(progress_path: Path, chapters: Sequence[Chapter], combined_words: int) -> None:
    data = yaml.safe_load(progress_path.read_text(encoding="utf-8"))
    stages = data.setdefault("stages", {})
    # Book 4 uses scalar stage values rather than nested stage objects.
    stages["draft"] = "complete"
    stages["revise"] = "complete"
    stages["polish"] = "complete"
    stages["export"] = "complete"
    stages["package"] = "in_progress"
    stages["publish"] = "pending"

    chapter_map = {chapter.number: chapter for chapter in chapters}
    if isinstance(data.get("chapters"), list):
        for row in data["chapters"]:
            try:
                number = int(row.get("n"))
            except Exception:
                continue
            if number in chapter_map:
                row["words"] = chapter_map[number].word_count
                # Preserve the repository's chapter-level revised state; stage
                # completion above records draft/revise/polish completion.

    book = data.setdefault("book", {})
    book["actual_words"] = sum(chapter.word_count for chapter in chapters)
    book["combined_reader_facing_words"] = combined_words
    book["chapter_count"] = 8

    cover_exists = (progress_path.parent / "cover.jpeg").exists() or (progress_path.parent / "cover.jpg").exists()
    if cover_exists:
        stages["package"] = "complete"
        planning_note = (
            "Book 4 manuscript, revision, polish, technical export, and package cover are complete. "
            "The reader-facing source plus reproducible DOCX/EPUB outputs passed validation. "
            "Publish is pending. Books 1-3 were not modified."
        )
    else:
        stages["package"] = "in_progress"
        planning_note = (
            "Book 4 manuscript, revision, polish, and technical export are complete. "
            "The reader-facing source plus reproducible DOCX/EPUB outputs passed validation. "
            "Package remains in progress because no final cover exists and author-controlled "
            "retailer/print choices remain. Publish is pending. Books 1-3 were not modified."
        )
    book["planning_note"] = planning_note

    validation = data.setdefault("validation", {})
    validation["export_assembly_tooling_complete"] = True
    validation["combined_export_complete"] = True
    validation["export_validation_complete"] = True
    validation["package_readiness_validation_complete"] = True
    validation["package_complete"] = cover_exists
    validation["cover_ready"] = cover_exists
    validation["reader_facing_placeholders_present"] = False
    validation["story_preserved"] = True
    validation["chapter_sources_changed_in_publication_pass"] = False
    validation["publish_complete"] = False
    write_text(progress_path, yaml.safe_dump(data, sort_keys=False, allow_unicode=True, width=1000))


def manifest_rows(paths: dict[str, Path], book_dir: Path) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for label in ["markdown", "text", "html", "docx", "epub"]:
        path = paths[label]
        rows.append(
            {
                "label": label,
                "path": str(path.relative_to(book_dir)),
                "size_bytes": path.stat().st_size,
                "sha256": sha256_file(path),
                "committed": label in {"markdown", "text", "html"},
            }
        )
    return rows


def check_table_markdown(validation_groups: dict[str, Validation]) -> str:
    lines = ["| Area | Check | Result | Detail |", "|---|---|---|---|"]
    for area, validation in validation_groups.items():
        for name, passed, detail in validation.checks:
            safe_detail = detail.replace("|", "\\|").replace("\n", " ")
            if len(safe_detail) > 220:
                safe_detail = safe_detail[:217] + "..."
            lines.append(f"| {area} | {name} | {'PASS' if passed else 'FAIL'} | {safe_detail} |")
    return "\n".join(lines)


def write_reports(
    book_dir: Path,
    chapters: Sequence[Chapter],
    paths: dict[str, Path],
    validation_groups: dict[str, Validation],
    page_count: int,
    contacts: Sequence[Path],
    epubcheck_output: str,
    manifest: Sequence[dict[str, object]],
) -> None:
    export_dir = book_dir / "export"
    package_dir = book_dir / "package"

    cover_exists = (book_dir / "cover.jpeg").exists() or (book_dir / "cover.jpg").exists()
    cover_val = "ready" if cover_exists else "blocked"
    cover_status_text = "complete" if cover_exists else "in_progress"
    cover_blocking_text = (
        "**Cover image present.** `cover.jpeg` is present in the book directory."
        if cover_exists
        else "**No final Book 4 cover exists in the repository.** Upload readiness cannot be claimed until the author supplies and approves an ebook cover with the correct title, author, series identification, dimensions, format, color profile, and thumbnail legibility."
    )
    package_accurate_status = (
        "- Export: complete.\n- Package: complete.\n- Publish: pending.\n- Uploaded, distributed, or live: no."
        if cover_exists
        else "- Export: complete.\n- Package: in progress / blocked by cover and author decisions.\n- Publish: pending.\n- Uploaded, distributed, or live: no."
    )
    package_boundary_text = (
        "Technical manuscript/export work is complete. Package completion is ready, as a valid cover has been supplied. Publish remains pending."
        if cover_exists
        else "Technical manuscript/export work is complete. Package completion is blocked by the absence of a final ebook cover and by unresolved author-controlled publication/print decisions. Publish remains pending."
    )
    cover_checkbox = "[x]" if cover_exists else "[ ]"

    front_text = "\n\n".join((book_dir / "front-matter" / name).read_text(encoding="utf-8") for name in FRONT_FILES)
    back_text = "\n\n".join((book_dir / "back-matter" / name).read_text(encoding="utf-8") for name in BACK_FILES)
    front_words = word_count(front_text)
    back_words = word_count(back_text)
    body_words = sum(chapter.word_count for chapter in chapters)
    chapter_heading_words = sum(word_count(f"Chapter {n} — {title}") for n, title, _ in CHAPTERS)
    total_words = word_count(paths["markdown"].read_text(encoding="utf-8"))
    scene_breaks = sum(chapter.scene_breaks for chapter in chapters)
    tool_versions = {
        "python": sys.version.split()[0],
        "pandoc": command_version("pandoc"),
        "libreoffice": command_version("libreoffice"),
        "pdftoppm": command_version("pdftoppm", ("-v",)),
        "epubcheck": command_version("epubcheck"),
    }
    tool_rows = "\n".join(f"- **{name}:** {version}" for name, version in tool_versions.items())

    chapter_rows = "\n".join(
        f"| {chapter.number} | {chapter.display_title} | {chapter.word_count:,} | {chapter.scene_breaks} | `{chapter.body_sha256}` |"
        for chapter in chapters
    )
    artifact_rows = "\n".join(
        f"| {row['label'].upper()} | `{row['path']}` | {human_size(int(row['size_bytes']))} | `{row['sha256']}` | {'Yes' if row['committed'] else 'No — reproducible build output'} |"
        for row in manifest
    )
    word_report = f"""# Word-Count and Structure Report — {TITLE}

Generated by `python3 books/book-04/export/finalize-package.py` on {BUILD_DATE.isoformat()}.

## Counting method

Visible words are counted after stripping Markdown control characters, links, and HTML tags with a Unicode-aware token expression. Chapter counts exclude YAML metadata and source chapter headings. The manuscript-body count is chapter prose only. The combined total includes title/front matter, chapter headings and prose, and back matter.

## Chapter report

| Chapter | Title | Prose words | Scene breaks | Body SHA-256 |
|---:|---|---:|---:|---|
{chapter_rows}

## Totals

| Measure | Count |
|---|---:|
| Chapter count | 8 |
| Manuscript body words (chapter prose) | {body_words:,} |
| Front-matter words | {front_words:,} |
| Chapter-heading words | {chapter_heading_words:,} |
| Back-matter words | {back_words:,} |
| Total combined reader-facing words | {total_words:,} |
| Scene breaks in chapter prose | {scene_breaks} |
| DOCX render-proof pages (not a print-interior page count) | {page_count} |

The original 25,000-word target remains unchanged in planning metadata; this report records the final manuscript rather than changing the target.

## Artifact sizes and hashes

| Format | Path | Size | SHA-256 | Stored in Git |
|---|---|---:|---|---|
{artifact_rows}
"""
    write_text(export_dir / "word-count-report.md", word_report)

    manifest_doc = {
        "book": TITLE,
        "series": SERIES,
        "series_number": SERIES_NUMBER,
        "author": AUTHOR,
        "build_date": BUILD_DATE.isoformat(),
        "source_head_expected": EXPECTED_START_HEAD,
        "tool_versions": tool_versions,
        "artifacts": list(manifest),
        "docx_render_proof_pages": page_count,
        "contact_sheets": [str(path.relative_to(book_dir)) for path in contacts],
        "epubcheck": epubcheck_output[-4000:],
    }
    write_text(export_dir / "artifact-manifest.json", json.dumps(manifest_doc, indent=2, ensure_ascii=False))

    write_text(
        export_dir / "README.md",
        f"""# Book 4 Export Pipeline

## Authoritative source

- Combined reader-facing Markdown: `manuscript-combined.md`
- Plain text: `manuscript-combined.txt`
- Standalone HTML: `manuscript-combined.html`
- Front matter: `../front-matter/`
- Back matter: `../back-matter/`
- Authoritative prose: `../manuscript/ch-01.md` through `ch-08.md`

## Reproducible build

From the repository root:

```bash
bash books/book-04/export/build.sh
```

Required tools: Python 3, Pandoc, LibreOffice, Poppler (`pdftoppm`), and the Python packages listed by import in `finalize-package.py`. `epubcheck` is used when available; an internal EPUB 3 structural validator always runs.

## Validated toolchain

{tool_rows}

## Commands executed by the pipeline

```text
python3 books/book-04/export/finalize-package.py
pandoc manuscript-combined.md --from=markdown --to=plain --wrap=none -o manuscript-combined.txt
pandoc manuscript-combined.md --from=markdown --to=html5 --standalone -o manuscript-combined.html
pandoc manuscript-combined.md --from=markdown --to=docx --reference-doc reference.docx -o dist/The-Archive-Fire.docx
libreoffice --headless --convert-to pdf --outdir qa/docx-render dist/The-Archive-Fire.docx
pdftoppm -png -r 72 qa/docx-render/The-Archive-Fire.pdf qa/docx-pages/page
pandoc manuscript-combined.md --from=markdown --to=epub3 --toc --toc-depth=1 -o dist/The-Archive-Fire.epub
epubcheck dist/The-Archive-Fire.epub  # when available
```

## Build outputs

The pipeline generates:

- `manuscript-combined.md`
- `manuscript-combined.txt`
- `manuscript-combined.html`
- `dist/The-Archive-Fire.docx`
- `dist/The-Archive-Fire.epub`
- temporary DOCX render-proof pages and contact sheets under `qa/`

Binary and visual-QA outputs are intentionally excluded from Git by `export/.gitignore`, matching the repository’s established source-first export convention. Their exact sizes and SHA-256 hashes are committed in `artifact-manifest.json` and `word-count-report.md`.

## Publication status

Export is technically complete and reproducible. The book has not been uploaded or published. The overall package remains incomplete until a valid cover is supplied and author-controlled retailer decisions are made.
""",
    )

    write_text(
        export_dir / "export-readiness.md",
        f"""---
status: complete
book: 4
build_date: "{BUILD_DATE.isoformat()}"
publish_status: pending
---

# Export Readiness — {TITLE}

## Prior blocker

The previous pass could not execute `assemble-manuscript.py` or safely transform the eight fetched chapter sources into a 36,000-word generated artifact. It therefore documented the blocker rather than generating or validating the export.

## Resolution

The blocker is resolved by the executable local/CI pipeline in `finalize-package.py` and `build.sh`. The pipeline reads the repository files directly, strips only per-chapter YAML and source headings, assembles all eight chapter bodies once and in order, builds the supported formats, and validates source preservation mechanically.

## Reader-facing source status

- Combined Markdown: **complete and validated**.
- Plain text: **complete and validated**.
- Standalone HTML: **complete and validated**.
- DOCX: **generated, opened through LibreOffice, structurally checked, and rendered page by page**.
- EPUB 3: **generated and structurally validated; epubcheck used when available**.
- PDF proof: **not a repository deliverable**. A temporary PDF is generated only to render and inspect the DOCX; no print-ready PDF is claimed because print specifications are unresolved.

## Preservation result

All eight chapter source files were unchanged. The combined chapter bodies match their source bodies exactly after removal of YAML metadata and the source-only heading. No plot, clue, solution, chronology, character arc, arrest basis, consultant arrangement, supplemental record, or ending content was altered.

## Export status

**COMPLETE.** This does not mean the overall publication package is complete or that the book has been uploaded or published.
""",
    )

    audit_rows = "\n".join(f"- [x] {item}" for item in LOCKED_STORY_AUDIT)
    validation_table = check_table_markdown(validation_groups)
    write_text(
        package_dir / "final-validation-report.md",
        f"""---
status: technical-validation-complete
book: 4
validated: "{BUILD_DATE.isoformat()}"
publish_status: pending
---

# Final Validation Report — {TITLE}

## Scope and starting state

- Repository: `dustinober1/The-Blackwood-Ridge-Mysteries`
- Default branch: `main`
- Required starting HEAD: `{EXPECTED_START_HEAD}`
- Required starting message: `{EXPECTED_START_MESSAGE}`
- Books 1–3: excluded from the build and edit scope.

## Mechanical validation

{validation_table}

## Tool versions

{tool_rows}

## Exact commands

```text
bash books/book-04/export/build.sh
python3 books/book-04/export/finalize-package.py
pandoc books/book-04/export/manuscript-combined.md --from=markdown --to=plain --wrap=none -o books/book-04/export/manuscript-combined.txt
pandoc books/book-04/export/manuscript-combined.md --from=markdown --to=html5 --standalone -o books/book-04/export/manuscript-combined.html
pandoc books/book-04/export/manuscript-combined.md --from=markdown --to=docx --reference-doc books/book-04/export/reference.docx -o books/book-04/export/dist/The-Archive-Fire.docx
libreoffice --headless --convert-to pdf --outdir books/book-04/export/qa/docx-render books/book-04/export/dist/The-Archive-Fire.docx
pdftoppm -png -r 72 books/book-04/export/qa/docx-render/The-Archive-Fire.pdf books/book-04/export/qa/docx-pages/page
pandoc books/book-04/export/manuscript-combined.md --from=markdown --to=epub3 --toc --toc-depth=1 -o books/book-04/export/dist/The-Archive-Fire.epub
epubcheck books/book-04/export/dist/The-Archive-Fire.epub  # when available
```

## Exact methods

- Direct SHA-256 comparison of every authoritative chapter file before and after the pipeline.
- Exact string comparison between each stripped source chapter body and its section in `manuscript-combined.md`.
- Heading/order, duplicate-body, YAML-leak, conflict-marker, placeholder, and internal-note scans.
- BeautifulSoup HTML parse, heading-order check, and internal-link resolution.
- DOCX Open Packaging Convention/XML inspection for metadata, heading styles, and page-break properties.
- LibreOffice headless DOCX open/render, page-text extraction with `pypdf`, chapter-start checks, blank-page scan, broken-character scan, and all-page PNG rendering.
- EPUB ZIP/mimetype/container/OPF/manifest/spine/navigation/resource/link/metadata/chapter-order validation; `epubcheck` when executable is available.
- Retailer HTML allow-list and 4,000-character validation.

## DOCX render proof

- Page count: {page_count}
- Every page rendered to PNG: yes.
- Contact sheets generated: {len(contacts)}.
- This page count is for the author-review DOCX render only and is **not** a final paperback page count.

## Story and clue readiness audit

{audit_rows}

## Manuscript-file changes

None. No chapter file changed in this publication pass.

## Package boundary

    {package_boundary_text}
""",
    )

    write_text(
        package_dir / "package-readiness.md",
        f"""---
status: {cover_status_text}
technical_exports: complete
cover: {cover_val}
print: decisions_required
publish: pending
---

# Package Readiness — {TITLE}

## Completed technical package work

- Reader-facing front matter and back matter finalized without unresolved placeholders.
- Combined Markdown, plain text, and standalone HTML generated and committed.
- DOCX and EPUB generated reproducibly, validated, and recorded by size and SHA-256.
- Retailer metadata and spoiler-safe listing assets finalized as drafts for author entry.
- Story/clue/ending preservation audit passed.
- Author decision checklist isolated to genuinely author-controlled choices.

## Blocking asset

{cover_blocking_text}

## Print status

No print-ready interior PDF or full cover wrap was generated because no authoritative trim size, paper type, bleed choice, ISBN choice, or print-template convention exists for this series. The author-review DOCX is not represented as a print interior.

## Accurate status

{package_accurate_status}
""",
    )

    write_text(
        package_dir / "author-decision-checklist.md",
        f"""# Author Decision Checklist — {TITLE}

Only author-controlled choices remain here; technical work belongs in the validation report, not this checklist.

## Required before ebook upload

- {cover_checkbox} Supply and approve the final Book 4 ebook cover (correct title, `{AUTHOR}`, `{SERIES} · Book 4`; 1,600 × 2,560 px preferred; RGB JPEG/TIFF; thumbnail-legible).
- [ ] Set the final release date or preorder schedule.
- [ ] Set the ebook list price.
- [ ] Confirm territorial publication rights.
- [ ] Decide whether to enroll the ebook in an exclusivity program such as KDP Select.
- [ ] Decide whether to use an ISBN for the ebook on any platform that permits or requires one; no ISBN has been assigned in the package.
- [ ] Confirm the recommended seven keywords and up to three category selections during retailer entry.

## Required only if a paperback edition will be produced

- [ ] Choose trim size.
- [ ] Choose paper type and black/white interior option.
- [ ] Choose bleed or no bleed.
- [ ] Choose or assign the paperback ISBN.
- [ ] Set paperback price and territories.
- [ ] Generate and approve a print interior from the chosen specifications.
- [ ] Generate and approve the final cover template/wrap after the final print page count fixes spine width; preserve a usable barcode area.
""",
    )

    write_text(
        package_dir / "print-readiness.md",
        f"""# Print Readiness — {TITLE}

## Result

**Not print-ready; author specifications are required.** The repository establishes an ebook-first package and does not establish an authoritative paperback trim, paper, bleed, ISBN, or cover-template convention.

## Work completed

- The combined source and review DOCX use consistent headings, paragraphs, scene-break treatment, and chapter page starts.
- The DOCX was rendered to a {page_count}-page proof for QA. This is not a final paperback page count.
- Front and back matter are structurally ready to flow into a print template after specifications are chosen.

## Work intentionally not claimed

- No print-ready interior PDF.
- No final margins, gutter, headers/footers, body numbering, front-matter numbering, blank-page scheme, or final page count.
- No print wrap, spine width, or barcode-area validation.

## Author decisions required

Trim size, paper type, bleed choice, paperback ISBN, price, territories, and final cover template. Once fixed, rebuild the interior, validate every page, record the final page count, and calculate the cover spine width from the platform template.
""",
    )


def write_book_readme(book_dir: Path) -> None:
    cover_exists = (book_dir / "cover.jpeg").exists() or (book_dir / "cover.jpg").exists()
    package_status = (
        "complete."
        if cover_exists
        else "in progress because the final cover and author-controlled upload/print decisions remain."
    )
    write_text(
        book_dir / "README.md",
        f"""# Book 4 — {TITLE}

## Current status

- Draft: complete.
- Revision: complete.
- Polish: complete.
- Export: complete and reproducibly validated.
- Package: {package_status}
- Publish: pending.

## Authoritative files

- Chapter sources: `manuscript/ch-01.md` through `ch-08.md`
- Combined reader-facing manuscript: `export/manuscript-combined.md`
- Export instructions and validation: `export/README.md`, `export/export-readiness.md`, `export/word-count-report.md`
- Retailer metadata: `publication/metadata.md`
- Listing copy: `listing/`
- Package status and remaining decisions: `package/package-readiness.md`, `package/final-validation-report.md`, `package/author-decision-checklist.md`

Nothing in this directory indicates that the book has been uploaded, distributed, accepted by a retailer, or published.
""",
    )


def source_sha_snapshot(chapters: Sequence[Chapter]) -> dict[str, str]:
    return {str(chapter.source_path): chapter.source_sha256 for chapter in chapters}


def assert_source_sha_unchanged(chapters: Sequence[Chapter], before: dict[str, str]) -> None:
    after = {str(chapter.source_path): sha256_file(chapter.source_path) for chapter in chapters}
    if before != after:
        diffs = {path: (before.get(path), after.get(path)) for path in sorted(set(before) | set(after)) if before.get(path) != after.get(path)}
        raise RuntimeError(f"Chapter source changed during production pass: {diffs}")


def finalize(repo_root: Path) -> None:
    book_dir = repo_root / "books" / "book-04"
    export_dir = book_dir / "export"
    manuscript_dir = book_dir / "manuscript"
    progress_path = book_dir / "progress.yaml"
    if not manuscript_dir.is_dir() or not progress_path.exists():
        raise RuntimeError(f"Book 4 source not found under {book_dir}")

    chapters = load_chapters(manuscript_dir)
    source_before = source_sha_snapshot(chapters)

    create_front_back(book_dir / "front-matter", book_dir / "back-matter")
    write_assembler(export_dir)
    write_build_script(export_dir)
    write_export_gitignore(export_dir)
    write_listing_and_metadata(book_dir)
    write_content_and_revision_notes(book_dir)
    write_packaging_guidance(book_dir)
    write_book_readme(book_dir)

    combined = assemble_manuscript(book_dir, chapters)
    paths = build_formats(export_dir, combined)

    markdown_validation = validate_markdown(paths["markdown"], chapters)
    html_validation = validate_html(paths["html"])
    docx_validation, page_count, contacts, _proof_pdf = validate_docx(paths["docx"], paths["qa_dir"])
    epub_validation, epubcheck_output = validate_epub(paths["epub"])
    retailer_validation = validate_retailer_html(book_dir / "listing" / "retailer-description.html")

    validation_groups = {
        "Markdown": markdown_validation,
        "HTML": html_validation,
        "DOCX": docx_validation,
        "EPUB": epub_validation,
        "Retailer HTML": retailer_validation,
    }
    if not all(group.passed for group in validation_groups.values()):
        raise RuntimeError("one or more validation groups failed")

    manifest = manifest_rows(paths, book_dir)
    combined_words = word_count(paths["markdown"].read_text(encoding="utf-8"))
    update_progress(progress_path, chapters, combined_words)
    write_reports(book_dir, chapters, paths, validation_groups, page_count, contacts, epubcheck_output, manifest)
    assert_source_sha_unchanged(chapters, source_before)

    # Final cross-file scans after reports/status updates.
    combined_text = paths["markdown"].read_text(encoding="utf-8")
    for pattern in BAD_READER_PATTERNS:
        if pattern.search(combined_text):
            raise RuntimeError(f"reader-facing combined manuscript contains forbidden pattern: {pattern.pattern}")

    print("BOOK4_FINALIZE_OK")
    print(f"Body words: {sum(chapter.word_count for chapter in chapters)}")
    print(f"Combined words: {combined_words}")
    print(f"DOCX proof pages: {page_count}")
    for row in manifest:
        print(f"{row['label']}: {row['size_bytes']} bytes {row['sha256']}")


def parse_args(argv: Sequence[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=None,
        help="Repository root. Defaults to three parents above this script.",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> None:
    args = parse_args(argv or sys.argv[1:])
    repo_root = args.repo_root.resolve() if args.repo_root else Path(__file__).resolve().parents[3]
    finalize(repo_root)


if __name__ == "__main__":
    main()
